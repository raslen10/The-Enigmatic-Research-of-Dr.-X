import os
import time
import chromadb
from chromadb.utils import embedding_functions
import PyPDF2
import pandas as pd
import docx
import ollama
from typing import List, Dict, Tuple, Optional
import tiktoken
from rouge_score import rouge_scorer
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor
from langdetect import detect

# Configuration
DATA_FOLDER = "data"
CHROMA_DB_PATH = "chroma_db"
EMBEDDING_MODEL = "nomic-embed-text"
LLM_MODEL = "llama3"

@dataclass
class PerformanceMetrics:
    process: str
    tokens: int
    time_taken: float
    tokens_per_sec: float

class ResearchProcessor:
    def __init__(self):
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
        self.performance_data = []
        self._init_vector_db()
    
    def _init_vector_db(self):
        self.chroma_client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
        self.nomic_ef = embedding_functions.OllamaEmbeddingFunction(
            url="http://localhost:11434",
            model_name=EMBEDDING_MODEL
        )
        self.collection = self.chroma_client.get_or_create_collection(
            name="drx_research",
            embedding_function=self.nomic_ef
        )
    
    def _track_performance(self, process: str, input_tokens: int, output_tokens: int, time_taken: float):
        metrics = PerformanceMetrics(
            process=process,
            tokens=input_tokens + output_tokens,
            time_taken=time_taken,
            tokens_per_sec=(input_tokens + output_tokens)/time_taken if time_taken > 0 else 0
        )
        self.performance_data.append(metrics)
        return metrics
    
    def detect_language(self, text: str) -> str:
        try:
            return detect(text[:500])
        except:
            return "unknown"
    
    def process_pdf(self, file_path: str) -> Tuple[List[str], List[Dict]]:
        texts = []
        metadatas = []
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    texts.append(text)
                    metadatas.append({
                        "source": os.path.basename(file_path),
                        "page": page_num,
                        "type": "pdf",
                        "position": f"page_{page_num}",
                        "language": self.detect_language(text)
                    })
        return texts, metadatas
    
    def process_docx(self, file_path: str) -> Tuple[List[str], List[Dict]]:
        texts = []
        metadatas = []
        doc = docx.Document(file_path)
        for para_num, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                texts.append(para.text)
                metadatas.append({
                    "source": os.path.basename(file_path),
                    "paragraph": para_num,
                    "type": "docx",
                    "position": f"para_{para_num}",
                    "language": self.detect_language(para.text)
                })
        return texts, metadatas
    
    def process_excel(self, file_path: str) -> Tuple[List[str], List[Dict]]:
        texts = []
        metadatas = []
        
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
            sheet_name = "CSV"
        else:
            xls = pd.ExcelFile(file_path)
            df = pd.concat([pd.read_excel(xls, sheet) for sheet in xls.sheet_names])
        
        text = df.to_markdown(index=False)
        if text.strip():
            texts.append(text)
            metadatas.append({
                "source": os.path.basename(file_path),
                "type": "excel" if file_path.endswith('.xlsx') else "csv",
                "position": "full_content",
                "language": self.detect_language(text)
            })
        return texts, metadatas
    
    def process_file(self, file_path: str) -> Tuple[List[str], List[Dict]]:
        if file_path.endswith('.pdf'):
            return self.process_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self.process_docx(file_path)
        elif file_path.endswith(('.xlsx', '.xls', '.csv')):
            return self.process_excel(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                return [content], [{
                    "source": os.path.basename(file_path),
                    "type": "text",
                    "position": "full_content",
                    "language": self.detect_language(content)
                }]
    
    def chunk_text(self, texts: List[str], metadatas: List[Dict]) -> Tuple[List[str], List[Dict]]:
        """
        Break down texts into manageable chunks with comprehensive metadata.
        
        Args:
            texts: List of texts to be chunked
            metadatas: List of corresponding metadata dictionaries
            
        Returns:
            Tuple of (chunks, chunk_metadatas) where:
            - chunks: List of text chunks
            - chunk_metadatas: List of metadata dictionaries for each chunk
        """
        chunks = []
        chunk_metadatas = []
        global_chunk_counter = 1  # Track chunks across all documents
        
        for text, metadata in zip(texts, metadatas):
            # Encode the text into tokens
            tokens = self.tokenizer.encode(text)
            
            # Calculate the number of chunks needed
            total_tokens = len(tokens)
            num_chunks = (total_tokens // 500) + (1 if total_tokens % 500 != 0 else 0)
            
            for i in range(0, total_tokens, 500):
                # Get the current chunk of tokens
                chunk_tokens = tokens[i:i+500]
                
                # Decode back to text
                chunk_text = self.tokenizer.decode(chunk_tokens)
                
                # Calculate position information
                start_token = i
                end_token = min(i + 500, total_tokens) - 1
                token_percentage = (end_token + 1) / total_tokens * 100
                
                # Create comprehensive metadata
                chunk_metadata = {
                    "source": metadata.get("source", "unknown"),
                    "document_type": metadata.get("type", "unknown"),
                    "original_language": metadata.get("language", "unknown"),
                    "position_info": metadata.get("position", "unknown"),
                    "page_number": metadata.get("page", metadata.get("paragraph", "unknown")),
                    "chunk_number": global_chunk_counter,
                    "document_chunk_sequence": (i // 500) + 1,  # Chunk number within this document
                    "total_chunks_in_doc": num_chunks,
                    "token_count": len(chunk_tokens),
                    "start_token": start_token,
                    "end_token": end_token,
                    "token_percentage": f"{token_percentage:.1f}%",
                    "is_continuation": i > 0  # Whether this chunk continues from previous
                }
                
                # Add any additional metadata from the original
                chunk_metadata.update({
                    k: v for k, v in metadata.items() 
                    if k not in chunk_metadata and not k.startswith('_')
                })
                
                chunks.append(chunk_text)
                chunk_metadatas.append(chunk_metadata)
                global_chunk_counter += 1
        
        return chunks, chunk_metadatas


    def load_documents(self, data_folder: str = DATA_FOLDER):
        if not os.path.exists(data_folder):
            raise FileNotFoundError(f"Data folder '{data_folder}' not found")
        
        for root, _, files in os.walk(data_folder):
            for file in files:
                file_path = os.path.join(root, file)
                try:
                    texts, metadatas = self.process_file(file_path)
                    if not texts:
                        continue
                    
                    chunks, chunk_metadatas = self.chunk_text(texts, metadatas)
                    
                    with ThreadPoolExecutor(max_workers=4) as executor:
                        embeddings = list(executor.map(
                            lambda chunk: ollama.embeddings(
                                model=EMBEDDING_MODEL,
                                prompt=chunk
                            )['embedding'],
                            chunks
                        ))
                    
                    ids = [f"{os.path.basename(file_path)}-{i}" for i in range(len(chunks))]
                    self.collection.add(
                        documents=chunks,
                        embeddings=embeddings,
                        metadatas=chunk_metadatas,
                        ids=ids
                    )
                except Exception as e:
                    print(f"Error processing file {file_path}: {e}")
    
    def _retrieve_relevant_chunks(self, query: str, n_results: int = 5) -> Tuple[List[str], List[Dict]]:
        """Retrieve relevant chunks from the vector database"""
        query_embedding = ollama.embeddings(
            model=EMBEDDING_MODEL,
            prompt=query
        )['embedding']
        
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=n_results
        )
        
        if not results['documents']:
            return [], []
        
        return results['documents'][0], results['metadatas'][0]
    
    def translate_document(self, source_file: str, target_lang: str) -> Tuple[str, PerformanceMetrics]:
        """Translate an entire document using RAG approach"""
        start_time = time.time()
        
        # Retrieve all chunks for this document
        results = self.collection.get(
            where={"source": source_file},
            include=["documents", "metadatas"]
        )
        
        if not results['documents']:
            return "Document not found in database.", None
        
        # Build context and translate each chunk
        translated_chunks = []
        total_input_tokens = 0
        total_output_tokens = 0
        
        for doc, meta in zip(results['documents'], results['metadatas']):
            prompt = f"""Translate this text to {target_lang}. Preserve technical terms and formatting.
            Original language: {meta.get('language', 'unknown')}
            Text: {doc}"""
            
            response = ollama.generate(
                model=LLM_MODEL,
                prompt=prompt,
                options={'temperature': 0.1}
            )
            
            translated_chunks.append(response['response'])
            total_input_tokens += len(self.tokenizer.encode(prompt))
            total_output_tokens += len(self.tokenizer.encode(response['response']))
        
        # Combine translated chunks
        translated_text = "\n\n".join(translated_chunks)
        
        metrics = self._track_performance(
            f"translate_document_to_{target_lang}",
            total_input_tokens,
            total_output_tokens,
            time.time() - start_time
        )
        
        return translated_text, metrics
    
    def summarize_document(self, source_file: str, strategy: str = "abstractive") -> Tuple[str, dict, PerformanceMetrics]:
        """Summarize a document using RAG approach"""
        start_time = time.time()
        
        # Retrieve all chunks for this document
        results = self.collection.get(
            where={"source": source_file},
            include=["documents", "metadatas"]
        )
        
        if not results['documents']:
            return "Document not found in database.", {}, None
        
        # Build context for summarization
        context = "\n\n".join([
            f"Source: {meta['source']} [{meta['position']}]\nContent: {doc}"
            for doc, meta in zip(results['documents'], results['metadatas'])
        ])
        
        # Generate summary
        prompt = f"""Create a {strategy} summary of this document. Be concise and maintain key technical details.
        Document: {context}
        Summary:"""
        
        response = ollama.generate(
            model=LLM_MODEL,
            prompt=prompt,
            options={'temperature': 0.3}
        )
        summary = response['response']
        
        # Calculate ROUGE scores against the first few chunks (for evaluation)
        sample_reference = " ".join(results['documents'][:3])
        scores = self.scorer.score(sample_reference[:5000], summary)
        
        metrics = self._track_performance(
            f"summarize_document_{strategy}",
            len(self.tokenizer.encode(prompt)),
            len(self.tokenizer.encode(summary)),
            time.time() - start_time
        )
        
        return summary, scores, metrics
    
    def ask_question(self, question: str, target_lang: str = None) -> Tuple[str, List[Dict], dict]:
        """RAG-based Q&A with optional translation"""
        start_time = time.time()
        
        # Retrieve relevant chunks
        documents, metadatas = self._retrieve_relevant_chunks(question)
        if not documents:
            return "No relevant information found.", [], {}
        
        # Build context
        context = "\n\n".join([
            f"Source: {meta['source']} [{meta['position']}]\nContent: {doc}"
            for doc, meta in zip(documents, metadatas)
        ])
        
        # Generate answer
        prompt = f"""Answer the question using only the provided context. Be precise and cite sources.
        Context: {context}
        Question: {question}
        Answer:"""
        
        response = ollama.generate(
            model=LLM_MODEL,
            prompt=prompt,
            options={'temperature': 0.2}
        )
        answer = response['response']
        
        # Track RAG performance
        rag_metrics = self._track_performance(
            "rag_qa",
            len(self.tokenizer.encode(prompt)),
            len(self.tokenizer.encode(answer)),
            time.time() - start_time
        )
        
        # Translate if needed
        translation_metrics = None
        if target_lang:
            answer, translation_metrics = self._translate_text(answer, target_lang, context)
        
        # Prepare sources
        sources = [{
            "source": meta['source'],
            "position": meta['position'],
            "content": doc[:200] + "..."
        } for doc, meta in zip(documents, metadatas)]
        
        return answer, sources, {
            "rag": rag_metrics,
            "translation": translation_metrics
        }
    
    def _translate_text(self, text: str, target_lang: str, context: str = None) -> Tuple[str, PerformanceMetrics]:
        """Internal translation method with performance tracking"""
        start_time = time.time()
        prompt = f"""Translate to {target_lang}. Preserve formatting and technical terms.
        {f"Context: {context}" if context else ""}
        Text: {text}"""
        
        response = ollama.generate(
            model=LLM_MODEL,
            prompt=prompt,
            options={'temperature': 0.1}
        )
        translated = response['response']
        
        metrics = self._track_performance(
            f"translate_to_{target_lang}",
            len(self.tokenizer.encode(prompt)),
            len(self.tokenizer.encode(translated)),
            time.time() - start_time
        )
        
        return translated, metrics
    
    def get_performance_data(self):
        return self.performance_data

    def list_documents(self) -> List[Dict]:
        """List all unique documents in the database"""
        results = self.collection.get(include=["metadatas"])
        unique_sources = {}
        
        for meta in results['metadatas']:
            source = meta['source']
            if source not in unique_sources:
                unique_sources[source] = {
                    "type": meta['type'],
                    "language": meta.get('language', 'unknown'),
                    "chunks": 1
                }
            else:
                unique_sources[source]['chunks'] += 1
                
        return [{"source": k, **v} for k, v in unique_sources.items()]

    def extract_text_from_file(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return self._extract_text_from_docx(file_path)
        elif ext in [".xlsx", ".xls", ".csv"]:
            return self._extract_text_from_excel(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _extract_text_from_pdf(self, file_path: str) -> str:
        reader = PyPDF2.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    def _extract_text_from_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text

    def _extract_text_from_excel(self, file_path: str) -> str:
        df = pd.read_excel(file_path)
        return df.to_string(index=False)

    def summarize_text(self, text: str, strategy: str) -> Tuple[str, Dict]:
        if strategy == "abstractive":
            summary = self.embedding_model.summarize(text, method="abstractive")
        else:
            summary = self.embedding_model.summarize(text, method="extractive")

        scorer = rouge_scorer.RougeScorer(["rouge1", "rouge2", "rougeL"], use_stemmer=True)
        scores = scorer.score(text, summary)

        return summary, scores

    def translate_text(self, text: str, target_lang: str) -> str:
        detected_lang = detect(text)
        if detected_lang == target_lang:
            return text
        return self.embedding_model.translate(text, target_lang)

    def build_vector_db(self, chunks: List[str], metadatas: List[Dict]):
        embeddings = [self.nomic_ef.embed(chunk) for chunk in chunks]
        self.collection.add(
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadatas
        )

    def ask_question(self, question: str) -> str:
        question_embedding = self.embedding_model.embed(question)
        relevant_chunks = self.vector_db.query(question_embedding)
        return self.embedding_model.generate_answer(relevant_chunks)